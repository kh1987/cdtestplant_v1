from io import BytesIO
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.oxml.shape import CT_Picture
from docx.parts.image import ImagePart
from docx.text.run import Run
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# 路径工具
from utils.path_utils import project_path

def getParentRunNode(node):
    """传入oxml节点对象，获取其祖先节点的CT_R"""
    if isinstance(node, CT_R):
        return node
    return getParentRunNode(node.getparent())

def generate_temp_doc(doc_type: str, project_id: int, round_num=None):
    """ 该函数参数：
    :param round_num: 只有回归说明和回归记录有
    :param project_id: 项目id
    :param doc_type:大纲 sm:说明 jl:记录 bg:报告 hsm:回归测试说明 hjl:回归测试记录,默认路径为dg -> 所以如果传错就生成生成大纲了
    :return (to_tpl_file路径, seitai_final_file路径)
    """
    project_path_str = project_path(project_id)
    # 根据传入需要处理的文档类型，自动获路径
    prefix = Path.cwd() / 'media' / project_path_str
    template_file: Path = prefix / 'form_template' / 'products' / '测评大纲.docx'
    to_tpl_file: Path = prefix / 'temp' / '测评大纲.docx'
    seitai_final_file: Path = prefix / 'final_seitai' / '测评大纲.docx'
    if doc_type == 'sm':
        template_file = prefix / 'form_template' / 'products' / '测试说明.docx'
        to_tpl_file = prefix / 'temp' / '测试说明.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / '测试说明.docx'
    elif doc_type == 'jl':
        template_file = prefix / 'form_template' / 'products' / '测试记录.docx'
        to_tpl_file = prefix / 'temp' / '测试记录.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / '测试记录.docx'
    elif doc_type == 'bg':
        template_file = prefix / 'form_template' / 'products' / '测评报告.docx'
        to_tpl_file = prefix / 'temp' / '测评报告.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / '测评报告.docx'
    elif doc_type == 'hsm':
        template_file = prefix / 'form_template' / 'products' / '回归测试说明.docx'
        to_tpl_file = prefix / 'temp' / f'第{round_num}轮回归测试说明.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / f'第{round_num}轮回归测试说明.docx'
    elif doc_type == 'hjl':
        template_file = prefix / 'form_template' / 'products' / '回归测试记录.docx'
        to_tpl_file = prefix / 'temp' / f'第{round_num}轮回归测试记录.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / f'第{round_num}轮回归测试记录.docx'
    elif doc_type == 'wtd':
        template_file = prefix / 'form_template' / 'products' / '测试问题单.docx'
        to_tpl_file = prefix / 'temp' / '测试问题单.docx'
        seitai_final_file: Path = prefix / 'final_seitai' / '测试问题单.docx'
    # 定义找寻被复制文件根路径 - 后续会根据type找子路径
    output_files_path = prefix / 'output_dir'
    # 这里可能修改：TODO:现在是找不到的被拷贝的文件，去根路径找大纲的被拷贝文件，看后续有什么要求进行修改
    dg_copied_files = []  # TODO：以后看是否该这里
    exclusive_copied_files = []
    # 将被拷贝文件分别放入不同两个数组
    for file in output_files_path.iterdir():
        if file.is_file():
            if file.suffix == '.docx':
                dg_copied_files.append(file)
        elif file.is_dir():
            if file.stem == doc_type:  # 这里要求在output_dir中专属文件夹必须是sm/jl/hsm/hjl/bg/wtd不然无法生成
                for f in file.iterdir():
                    if f.suffix == '.docx':
                        exclusive_copied_files.append(f)
    # 找到基础模版的所有std域
    doc = Document(template_file.as_posix())
    body = doc.element.body
    sdt_element_list = body.xpath('./w:sdt')
    # 找到sdt域的名称 -> 为了对应output_dir文件 / 储存所有output_dir图片
    area_name_list = []
    image_part_list = []
    # 遍历所有控件 -> 放入area_name_list【这里准备提取公共代码】
    for sdt_ele in sdt_element_list:
        for elem in sdt_ele.iterchildren():
            # 获取“域”的名称
            if elem.tag.endswith('sdtPr'):
                for el in elem.getchildren():
                    if el.tag.endswith('alias'):
                        if len(el.attrib.values()) > 0:
                            area_name = el.attrib.values()[0]
                            area_name_list.append(area_name)
            # 开始替换里面的“域”
            if elem.tag.endswith('sdtContent'):
                elem.clear()
                if len(area_name_list) > 0:
                    area_pop_name = area_name_list.pop(0)
                    # 取到“域名称”，这里先去找media/output_dir/xx下文件，然后找media/output下文件
                    copied_file_path = ""
                    # TODO:还是那个问题，只区别了大纲 -> dg
                    if doc_type == 'dg':
                        for file in dg_copied_files:
                            if file.stem == area_pop_name:
                                copied_file_path = file
                    else:
                        if round_num is None:
                            for file in exclusive_copied_files:
                                if file.stem == area_pop_name:
                                    copied_file_path = file
                            # 这里判断是否copied_file_path没取到文件，然后遍历output_dir下文件
                            if not copied_file_path:
                                for file in dg_copied_files:
                                    if file.stem == area_pop_name:
                                        copied_file_path = file
                        else:
                            # 因为回归的轮次，前面会加 -> 第{round_num}轮
                            for file in exclusive_copied_files:  # 这里多了第{round_num}轮
                                if file.stem == f"第{round_num}轮{area_pop_name}":
                                    copied_file_path = file
                            if not copied_file_path:
                                for file in dg_copied_files:
                                    if file.stem == area_pop_name:
                                        copied_file_path = file
                    # 找到所需的文件，将其数据复制到对应area_name的“域”
                    if copied_file_path:
                        doc_copied = Document(copied_file_path)
                        copied_element_list = []
                        element_list = doc_copied.element.body.inner_content_elements
                        for elet in element_list:
                            if isinstance(elet, CT_P):
                                copied_element_list.append(Paragraph(elet, doc_copied))
                            if isinstance(elet, CT_Tbl):
                                copied_element_list.append(Table(elet, doc_copied))
                        for para_copied in copied_element_list:
                            elem.append(para_copied._element)

                        # 下面代码就是将图片全部提取到image_part_list，以便后续插入
                        doc_copied = Document(copied_file_path)  # 需要重新获取否则namespace错误
                        copied_body = doc_copied.element.body
                        img_node_list = copied_body.xpath('.//pic:pic')
                        if not img_node_list:
                            pass
                        else:
                            for img_node in img_node_list:
                                img: CT_Picture = img_node
                                # 根据节点找到图片的关联id
                                embed = img.xpath('.//a:blip/@r:embed')[0]
                                # 这里得到ImagePart -> 马上要给新文档添加
                                related_part: ImagePart = doc_copied.part.related_parts[embed]
                                # doc_copied.part.related_parts是一个字典
                                image_part_list.append(related_part)

    # 新文档添加sdt_element_list的图片
    graph_node_list = body.xpath('.//pic:pic')
    img_count = 0
    if len(graph_node_list) == len(image_part_list):
        for graph_node in graph_node_list:
            image_run_node = getParentRunNode(graph_node)
            image_run_node.clear()
            copied_bytes_io = BytesIO(image_part_list[img_count].image.blob)
            r_element = Run(image_run_node, doc)
            inline_shape = r_element.add_picture(copied_bytes_io)
            # 设置图片位置尺寸
            source_width = inline_shape.width
            inline_shape.width = Cm(12)
            inline_shape.height = int(inline_shape.height * (inline_shape.width / source_width))
            # 设置图片所在段落居中对齐
            r_element.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            img_count += 1
    else:
        return {'code': 'error', 'msg': '模版文件有自定义图片，请删除后重试...'}

    try:
        doc.save(to_tpl_file)
        return to_tpl_file, seitai_final_file
    except PermissionError as e:
        return {'code': 'error', 'msg': '生成的temp文件已打开，请关闭后重试...'}
