from pathlib import Path
from io import BytesIO
# 测试运行时间
import time

from ninja_extra.controllers import api_controller, ControllerBase, route
from ninja_jwt.authentication import JWTAuth
from ninja_extra.permissions import IsAuthenticated
from django.db import transaction
from django.shortcuts import get_object_or_404

# 文档处理相关库
from docxtpl import DocxTemplate
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.oxml.shape import CT_Picture
from docx.parts.image import ImagePart
from docx.text.run import Run
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 自己的工具模块
from utils.chen_response import ChenResponse

# 模型模块
from apps.project.models import Project

def getParentRunNode(node):
    """传入oxml节点对象，获取其祖先节点的CT_R"""
    if isinstance(node, CT_R):
        return node
    return getParentRunNode(node.getparent())

# @api_controller("/create", tags=['生成产品文档接口'], auth=JWTAuth(), permissions=[IsAuthenticated])
@api_controller("/create", tags=['生成产品文档接口'])
class GenerateSeitaiController(ControllerBase):
    @route.get("/dgDocument", url_name="create-dgDocument")
    @transaction.atomic
    def create_dgDocument(self, id: int):
        # 生成大纲所需output文件-系列文件-目录
        dg_series_documents_dir = Path.cwd() / 'media' / 'output_dir'
        # 大纲基本模版路径
        dg_seitai_path = Path.cwd() / 'media' / 'form_template' / 'products' / '测评大纲.docx'
        # 生成最终给docxtpl的模版路径
        dg_replace_path = Path.cwd() / 'media' / 'temp' / '测评大纲.docx'
        # 生成大纲最终文档路径
        dg_seitai_final_path = Path.cwd() / 'media' / 'final_seitai' / '测评大纲.docx'
        # 剔除所有的非docx文件
        copied_files = []
        for doc in dg_series_documents_dir.iterdir():
            if doc.suffix == '.docx':
                copied_files.append(doc)
        # 获取项目Model
        project_obj = get_object_or_404(Project, id=id)
        # 生成大纲需要的基础变量context
        context = {'is_JD': False}
        if project_obj.report_type == '9':
            context['is_JD'] = True
        context['ident'] = project_obj.ident
        # TODO:暂时声明为公开
        context['sec_title'] = '公开'
        context['sec'] = '公开'
        context['name'] = project_obj.name
        context['duty_person'] = project_obj.duty_person
        if len(project_obj.member) > 0:
            context['member'] = project_obj.member[0]
        else:
            context['member'] = context['duty_person']
        context['entrust_unit'] = project_obj.entrust_unit

        # 最后生成
        doc = Document(dg_seitai_path.as_posix())
        body = doc.element.body
        sdt_element_list = body.xpath('./w:sdt')
        # 储存找到的域的名称（即output_dir文件名）
        area_name_list = []
        # 储存找到所以的图片
        image_part_list = []
        # 遍历所以sdt控件 - 下面就是每个std控件
        for sdt_ele in sdt_element_list:
            for elem in sdt_ele.iterchildren():
                # 获取域的名称
                if elem.tag.endswith('sdtPr'):
                    for el in elem.getchildren():
                        if el.tag.endswith('alias'):
                            if len(el.attrib.values()) > 0:
                                area_name = el.attrib.values()[0]
                                area_name_list.append(area_name)
                if elem.tag.endswith('sdtContent'):
                    elem.clear()
                    if len(area_name_list) > 0:
                        kongjian_name = area_name_list.pop(0)
                        # 这里取到了-域名称
                        copied_file_path = ""
                        for file in copied_files:
                            if file.stem == kongjian_name:
                                copied_file_path = file
                        if copied_file_path:
                            doc_copied = Document(copied_file_path)
                            copied_element_list = []
                            # 被拷贝元素列表
                            element_list = doc_copied.element.body.inner_content_elements
                            for elet in element_list:
                                if isinstance(elet, CT_P):
                                    copied_element_list.append(Paragraph(elet, doc_copied))
                                if isinstance(elet, CT_Tbl):
                                    copied_element_list.append(Table(elet, doc_copied))
                            for para_copied in copied_element_list:
                                elem.append(para_copied._element)

                            # 去找图片放入parts图片的列表以便后续替换
                            doc_copied = Document(copied_file_path)
                            copied_body = doc_copied.element.body
                            img_node_list = copied_body.xpath('.//pic:pic')
                            if not img_node_list:
                                pass
                            else:
                                for img_node in img_node_list:
                                    img: CT_Picture = img_node
                                    # 根据节点找到图片的关联id
                                    embed = img.xpath('.//a:blip/@r:embed')[0]
                                    # 这里得到ImagePart -> 马上要给新文档添加
                                    related_part: ImagePart = doc_copied.part.related_parts[embed]
                                    # doc_copied.part.related_parts是一个字典
                                    image_part_list.append(related_part)
        # 最后在新文档中添加图片
        graph_node_list = body.xpath('.//pic:pic')
        img_count = 0
        if len(graph_node_list) == len(image_part_list):
            for graph_node in graph_node_list:
                image_run_node = getParentRunNode(graph_node)
                image_run_node.clear()
                copied_bytes_io = BytesIO(image_part_list[img_count].image.blob)
                r_element = Run(image_run_node, doc)
                inline_shape = r_element.add_picture(copied_bytes_io)
                # 设置图片位置尺寸
                source_width = inline_shape.width
                inline_shape.width = Cm(12)
                inline_shape.height = int(inline_shape.height * (inline_shape.width / source_width))
                # 设置图片所在段落居中对齐
                r_element.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                img_count += 1
        else:
            return ChenResponse(code=400, status=400, message='注意模版里面有自定义图片，请删除后重试!!!')

        try:
            doc.save(dg_replace_path)
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="注意你打开了生成的文档，请关闭后再试，{0}".format(e))

        doc = DocxTemplate(dg_replace_path)
        start = time.time()
        doc.render(context)  # 耗时最长，TODO:异步任务处理？或前端等待？
        end = time.time()
        print('渲染耗时：',end - start)
        try:
            doc.save(dg_seitai_final_path)
            return ChenResponse(status=200, code=200, message="最终大纲生成成功！")
        except PermissionError as e:
            return ChenResponse(status=400, code=400, message="模版文件已打开，请关闭后再试，{0}".format(e))
